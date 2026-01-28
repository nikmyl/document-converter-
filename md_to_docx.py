#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts Markdown (.md) files to properly formatted Word documents (.docx)
"""

import re
import sys
import argparse
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)

try:
    import markdown
    from markdown.extensions import tables, fenced_code, codehilite
except ImportError:
    print("Error: markdown is required. Install it with: pip install markdown")
    sys.exit(1)


class MarkdownToDocxConverter:
    """Convert Markdown files to formatted Word documents"""

    # Supported input extensions
    SUPPORTED_EXTENSIONS = {'.md', '.markdown', '.txt'}

    def __init__(self, input_file: str, output_file: Optional[str] = None):
        """
        Initialize the converter

        Args:
            input_file: Path to the input .md, .markdown, or .txt file
            output_file: Path to the output .docx file (optional)
        """
        self.input_file = Path(input_file)

        if not self.input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")

        if self.input_file.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError("Input file must be a .md, .markdown, or .txt file")

        if output_file:
            self.output_file = Path(output_file)
        else:
            self.output_file = self.input_file.with_suffix('.docx')

        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up custom styles for the document"""
        styles = self.doc.styles

        # Code block style
        try:
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            # Style already exists
            pass

        # Inline code style
        try:
            inline_code = styles.add_style('InlineCode', WD_STYLE_TYPE.CHARACTER)
            inline_code.font.name = 'Courier New'
            inline_code.font.size = Pt(10)
            inline_code.font.color.rgb = RGBColor(199, 37, 78)
        except ValueError:
            pass

    def convert(self) -> str:
        """
        Convert the Markdown file to DOCX

        Returns:
            Path to the output file
        """
        # Read the markdown content
        with open(self.input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Process the markdown line by line for better control
        self._process_markdown(md_content)

        # Save the document
        self.doc.save(str(self.output_file))
        return str(self.output_file)

    def _process_markdown(self, content: str):
        """Process markdown content and add to document"""
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_block_lines = []
        in_list = False
        list_items = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_lines = []
                else:
                    # End of code block
                    self._add_code_block('\n'.join(code_block_lines))
                    in_code_block = False
                    code_block_lines = []
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Handle headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                self._add_heading(text, level)
                i += 1
                continue

            # Handle horizontal rules
            if re.match(r'^(\*{3,}|-{3,}|_{3,})$', line.strip()):
                self.doc.add_paragraph('_' * 50)
                i += 1
                continue

            # Handle unordered lists
            list_match = re.match(r'^[\*\-\+]\s+(.+)$', line)
            if list_match:
                text = list_match.group(1)
                self._add_bullet(text)
                i += 1
                continue

            # Handle ordered lists
            ordered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
            if ordered_match:
                text = ordered_match.group(2)
                self._add_numbered_list_item(text)
                i += 1
                continue

            # Handle blockquotes
            if line.strip().startswith('>'):
                quote_text = line.strip()[1:].strip()
                self._add_blockquote(quote_text)
                i += 1
                continue

            # Handle tables
            if '|' in line and line.strip().startswith('|'):
                # Collect all table rows
                table_rows = []
                while i < len(lines) and '|' in lines[i]:
                    table_rows.append(lines[i])
                    i += 1

                if table_rows:
                    self._add_table(table_rows)
                continue

            # Handle empty lines
            if not line.strip():
                # Add a blank line only if the last paragraph isn't already empty
                if self.doc.paragraphs and self.doc.paragraphs[-1].text.strip():
                    self.doc.add_paragraph()
                i += 1
                continue

            # Handle regular paragraphs with inline formatting
            self._add_formatted_paragraph(line)
            i += 1

    def _add_heading(self, text: str, level: int):
        """Add a heading to the document"""
        # Clean the text from markdown formatting
        text = self._clean_text(text)
        heading_style = f'Heading {min(level, 9)}'
        self.doc.add_heading(text, level=level)

    def _add_formatted_paragraph(self, text: str):
        """Add a paragraph with inline formatting (bold, italic, code, links)"""
        paragraph = self.doc.add_paragraph()

        # Process inline formatting
        self._process_inline_formatting(paragraph, text)

    def _add_hyperlink(self, paragraph, text: str, url: str):
        """
        Add a clickable hyperlink to a paragraph

        Args:
            paragraph: The paragraph to add the link to
            text: The display text for the link
            url: The URL to link to
        """
        # Get the document part
        part = self.doc.part

        # Create relationship for the hyperlink
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run for the hyperlink text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Add blue color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)

        # Add underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        # Add the text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _process_inline_formatting(self, paragraph, text: str):
        """Process inline markdown formatting"""
        # Tokenize the text to handle formatting more efficiently
        # This avoids character-by-character processing

        pos = 0
        buffer = []  # Collect regular text

        def flush_buffer():
            """Add accumulated text to paragraph"""
            if buffer:
                paragraph.add_run(''.join(buffer))
                buffer.clear()

        while pos < len(text):
            # Check for inline code `code`
            if text[pos:pos+1] == '`':
                end = text.find('`', pos + 1)
                if end != -1:
                    flush_buffer()
                    code_text = text[pos+1:end]
                    run = paragraph.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(199, 37, 78)
                    pos = end + 1
                    continue

            # Check for bold+italic ***text*** or ___text___
            if text[pos:pos+3] in ('***', '___'):
                delimiter = text[pos:pos+3]
                end = text.find(delimiter, pos + 3)
                if end != -1:
                    flush_buffer()
                    inner_text = text[pos+3:end]
                    run = paragraph.add_run(inner_text)
                    run.bold = True
                    run.italic = True
                    pos = end + 3
                    continue

            # Check for bold **text** or __text__
            if text[pos:pos+2] in ('**', '__'):
                delimiter = text[pos:pos+2]
                end = text.find(delimiter, pos + 2)
                if end != -1:
                    flush_buffer()
                    bold_text = text[pos+2:end]
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                    pos = end + 2
                    continue

            # Check for italic *text* or _text_ (but not part of ** or __)
            if text[pos:pos+1] in ('*', '_'):
                # Make sure it's not part of ** or __
                if pos > 0 and text[pos-1] == text[pos]:
                    buffer.append(text[pos])
                    pos += 1
                    continue
                if pos + 1 < len(text) and text[pos+1] == text[pos]:
                    # This is start of ** or __, skip
                    pass
                else:
                    delimiter = text[pos:pos+1]
                    end = text.find(delimiter, pos + 1)
                    if end != -1 and end > pos + 1:
                        # Check it's not part of ** or __
                        if end + 1 >= len(text) or text[end+1] != delimiter:
                            flush_buffer()
                            italic_text = text[pos+1:end]
                            run = paragraph.add_run(italic_text)
                            run.italic = True
                            pos = end + 1
                            continue

            # Check for links [text](url)
            if text[pos:pos+1] == '[':
                close_bracket = text.find(']', pos)
                if close_bracket != -1 and close_bracket + 1 < len(text) and text[close_bracket + 1] == '(':
                    close_paren = text.find(')', close_bracket + 2)
                    if close_paren != -1:
                        flush_buffer()
                        link_text = text[pos+1:close_bracket]
                        link_url = text[close_bracket+2:close_paren]
                        # Add actual clickable hyperlink
                        self._add_hyperlink(paragraph, link_text, link_url)
                        pos = close_paren + 1
                        continue

            # Regular text - add to buffer
            buffer.append(text[pos])
            pos += 1

        # Flush any remaining text
        flush_buffer()

    def _add_bullet(self, text: str):
        """Add a bullet point"""
        text = self._clean_text(text)
        paragraph = self.doc.add_paragraph(style='List Bullet')
        self._process_inline_formatting(paragraph, text)

    def _add_numbered_list_item(self, text: str):
        """Add a numbered list item"""
        text = self._clean_text(text)
        paragraph = self.doc.add_paragraph(style='List Number')
        self._process_inline_formatting(paragraph, text)

    def _add_blockquote(self, text: str):
        """Add a blockquote"""
        text = self._clean_text(text)
        paragraph = self.doc.add_paragraph()
        paragraph.style = 'Quote'
        self._process_inline_formatting(paragraph, text)

    def _add_code_block(self, code: str):
        """Add a code block"""
        paragraph = self.doc.add_paragraph(code)
        paragraph.style = 'CodeBlock'

    def _add_table(self, table_rows: list):
        """Add a table to the document"""
        if not table_rows:
            return

        # Parse table rows into cells
        parsed_rows = []
        for row in table_rows:
            # Split by | and clean up
            cells = [cell.strip() for cell in row.split('|')]

            # Remove empty first/last elements (from leading/trailing |)
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]

            # Filter out the separator row (contains only - and :)
            if cells and not all(set(cell.replace('-', '').replace(':', '').replace(' ', '')) == set() for cell in cells):
                parsed_rows.append(cells)

        if not parsed_rows:
            return

        # Determine table dimensions
        num_cols = max(len(row) for row in parsed_rows)
        num_rows = len(parsed_rows)

        # Create the table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Populate the table
        for i, row_data in enumerate(parsed_rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                # Add text with inline formatting
                if cell_text:
                    # Create a paragraph in the cell
                    cell.text = ''  # Clear default text
                    paragraph = cell.paragraphs[0]
                    self._process_inline_formatting(paragraph, cell_text)

                    # Make header row bold
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

        # Add spacing after table
        self.doc.add_paragraph()

    def _clean_text(self, text: str) -> str:
        """Clean text from simple markdown artifacts"""
        # This is basic cleaning; inline formatting is handled separately
        return text.strip()


def main():
    """Main entry point for the script"""
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Word documents (.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s input.md
  %(prog)s input.md -o output.docx
  %(prog)s *.md
        """
    )

    parser.add_argument(
        'input_files',
        nargs='+',
        help='Input Markdown file(s) to convert'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output file path (only valid with single input file)'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output'
    )

    args = parser.parse_args()

    # Validate arguments
    if args.output and len(args.input_files) > 1:
        print("Error: -o/--output option can only be used with a single input file")
        sys.exit(1)

    # Process each file
    for input_file in args.input_files:
        try:
            if args.verbose:
                print(f"Converting: {input_file}")

            converter = MarkdownToDocxConverter(
                input_file,
                args.output if args.output else None
            )
            output_path = converter.convert()

            print(f"[OK] Converted: {input_file} -> {output_path}")

        except Exception as e:
            print(f"[ERROR] Error converting {input_file}: {str(e)}", file=sys.stderr)
            if args.verbose:
                import traceback
                traceback.print_exc()


if __name__ == '__main__':
    main()
